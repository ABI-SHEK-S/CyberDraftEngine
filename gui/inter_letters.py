from pathlib import Path
import sys
import tkinter as tk
from tkinter import ttk, messagebox
import os
from docx import Document
from db.database import save_case
from datetime import datetime
import re
import logging

try:
    from tkcalendar import DateEntry
    TKCALENDAR_AVAILABLE = True
except ImportError:
    TKCALENDAR_AVAILABLE = False

from .utils import replace_placeholder_in_paragraph


class InterLetters:
    def __init__(self, parent, app):
        self.parent = parent
        self.app = app
        self.entries = []
        self.entry_frames = []
        self.platform_type = None
        self.google_id_type = tk.StringVar(value="Gmail ID")

        # Logging setup
        log_dir = os.path.join(Path.home(), 'Documents', 'LetterGeneratorLogs')
        os.makedirs(log_dir, exist_ok=True)
        logging.basicConfig(
            filename=os.path.join(log_dir, 'inter_letters.log'),
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )

        self.inter_template_dir = os.path.join(self.app.template_dir, 'inter') if self.app.template_dir else None
        self.setup_ui()

    # ---------------------------
    # UI Setup
    # ---------------------------
    def setup_ui(self):
        inter_inner = tk.Frame(self.parent, bg="white")
        inter_inner.pack(pady=20, padx=60, fill="both", expand=True)
        inter_inner.grid_columnconfigure(0, weight=1)
        inter_inner.grid_rowconfigure(2, weight=1)

        # Platform Selector
        tk.Label(inter_inner, text="Platform:", bg="white", font=("Segoe UI", 10, "bold"),
                 fg=self.app.text_color).grid(row=0, column=0, sticky="w", pady=(0, 5))

        self.inter_option = ttk.Combobox(
            inter_inner, values=["WhatsApp", "Facebook", "Instagram", "Google", "Twitter"],
            state="readonly"
        )
        self.inter_option.set("Select Platform")
        self.inter_option.grid(row=1, column=0, sticky="ew", pady=5)
        self.inter_option.bind("<<ComboboxSelected>>", self.on_platform_change)

        # Scrollable Field Area
        self.canvas = tk.Canvas(inter_inner, bg="white", highlightthickness=0)
        self.canvas.grid(row=2, column=0, sticky="nsew", pady=5)
        scrollbar = ttk.Scrollbar(inter_inner, orient="vertical", command=self.canvas.yview)
        scrollbar.grid(row=2, column=1, sticky="ns")
        self.canvas.configure(yscrollcommand=scrollbar.set)

        self.field_frame = tk.Frame(self.canvas, bg="white")
        self.window_id = self.canvas.create_window((0, 0), window=self.field_frame, anchor="nw")
        self.field_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.canvas.bind("<Configure>", self._resize_canvas_frame)

        # Mouse wheel: bind/unbind when cursor enters/leaves scroll area
        self.field_frame.bind("<Enter>", self._bind_mousewheel)
        self.field_frame.bind("<Leave>", self._unbind_mousewheel)

        # Add Field Button
        self.add_field_button = ttk.Button(inter_inner, text="+ Add Field", command=self.add_field)
        self.add_field_button.grid(row=3, column=0, sticky="w", pady=(0, 10))

        # Date Range
        self.inter_date_frame = tk.Frame(inter_inner, bg="white")
        self.inter_date_frame.grid(row=4, column=0, sticky="ew", pady=5)
        tk.Label(self.inter_date_frame, text="Date Range:", font=("Segoe UI", 10, "bold"),
                 bg="white", fg=self.app.text_color).pack(anchor="w", pady=(0, 2))
        date_inner = tk.Frame(self.inter_date_frame, bg="white")
        date_inner.pack(anchor="w")

        tk.Label(date_inner, text="From Date:", bg="white").pack(side="left", padx=5)
        if TKCALENDAR_AVAILABLE:
            self.inter_from_date_entry = DateEntry(date_inner, date_pattern='dd-mm-yyyy', width=12,
                                                   background=self.app.accent_color, foreground='white')
        else:
            self.inter_from_date_entry = ttk.Entry(date_inner, width=12)
        self.inter_from_date_entry.pack(side="left", padx=5)

        tk.Label(date_inner, text="To Date:", bg="white").pack(side="left", padx=5)
        if TKCALENDAR_AVAILABLE:
            self.inter_to_date_entry = DateEntry(date_inner, date_pattern='dd-mm-yyyy', width=12,
                                                 background=self.app.accent_color, foreground='white')
        else:
            self.inter_to_date_entry = ttk.Entry(date_inner, width=12)
        self.inter_to_date_entry.pack(side="left", padx=5)

        # Buttons
        button_frame = tk.Frame(inter_inner, bg="white")
        button_frame.grid(row=5, column=0, columnspan=2, pady=(10, 5))

        self.inter_generate_button = ttk.Button(
            button_frame, text="Generate Letter",
            command=self.generate_inter_letter, state="disabled", width=20
        )
        self.inter_generate_button.pack(pady=5)

        self.view_letters_inter_button = ttk.Button(
            button_frame, text="View Letters", command=self.view_letters_inter,
            state="disabled", width=20
        )
        self.view_letters_inter_button.pack(pady=5)

        # Status Label
        self.inter_status_label = tk.Label(inter_inner, text="", font=("Segoe UI", 10),
                                           bg="white", fg=self.app.success_color)
        self.inter_status_label.grid(row=6, column=0, sticky="w", pady=5)

    # ---------------------------
    # Scroll Helpers
    # ---------------------------
    def _resize_canvas_frame(self, event):
        self.canvas.itemconfigure(self.window_id, width=event.width)

    def _on_mousewheel(self, event):
        if event.num == 4:
            self.canvas.yview_scroll(-1, "units")
        elif event.num == 5:
            self.canvas.yview_scroll(1, "units")
        else:
            delta = -1 if event.delta > 0 else 1
            self.canvas.yview_scroll(delta, "units")

    def _bind_mousewheel(self, event):
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
        self.canvas.bind_all("<Button-4>", self._on_mousewheel)
        self.canvas.bind_all("<Button-5>", self._on_mousewheel)

    def _unbind_mousewheel(self, event):
        try:
            self.canvas.unbind_all("<MouseWheel>")
            self.canvas.unbind_all("<Button-4>")
            self.canvas.unbind_all("<Button-5>")
        except Exception:
            pass

    def auto_scroll_to_bottom(self):
        self.canvas.update_idletasks()
        self.canvas.yview_moveto(1)

    # ---------------------------
    # Label helper
    # ---------------------------
    def _label_for_index(self, idx):
        if self.platform_type == "WhatsApp":
            return f"WhatsApp account {idx}:"
        elif self.platform_type == "Google":
            return f"{self.google_id_type.get()} {idx}:"
        else:
            return f"URL {idx}:"

    # ---------------------------
    # Field Management
    # ---------------------------
    def add_field(self):
        # ✅ Prevent adding if no platform selected
        if not self.platform_type or self.platform_type == "Select Platform":
            messagebox.showwarning("Select Platform", "Please select a platform before adding fields.")
            return

        frame = tk.Frame(self.field_frame, bg="white")
        frame.grid(row=len(self.entries) + 1, column=0, sticky="ew", pady=5)
        frame.grid_columnconfigure(1, weight=1)

        label_text = self._label_for_index(len(self.entries) + 1)
        label = tk.Label(frame, text=label_text, bg="white", font=("Segoe UI", 10, "bold"))
        label.grid(row=0, column=0, padx=5, sticky="w")

        entry = ttk.Entry(frame)
        entry.grid(row=0, column=1, padx=5, sticky="ew")
        entry.bind("<KeyRelease>", self.update_generate_button_state)

        delete_btn = ttk.Button(frame, text="Delete", command=lambda: self.delete_field(frame))
        delete_btn.grid(row=0, column=2, padx=5)
        if len(self.entries) == 0:
            delete_btn.config(state="disabled")

        self.entries.append(entry)
        self.entry_frames.append(frame)

        self.auto_scroll_to_bottom()
        self.update_generate_button_state()


    def delete_field(self, frame):
        if len(self.entries) > 1:
            idx = self.entry_frames.index(frame)
            self.entries.pop(idx)
            self.entry_frames.remove(frame)
            frame.destroy()

            for i, fr in enumerate(self.entry_frames):
                fr.winfo_children()[0].config(text=self._label_for_index(i + 1))

            if len(self.entries) == 1:
                self.entry_frames[0].winfo_children()[2].config(state="disabled")

            self.update_generate_button_state()

    def show_google_id_selector(self):
        selector_frame = tk.Frame(self.field_frame, bg="white")
        selector_frame.grid(row=0, column=0, columnspan=3, pady=(0, 5), sticky="w")
        tk.Label(selector_frame, text="Google ID Type:", bg="white",
                 font=("Segoe UI", 10, "bold")).pack(side="left", padx=(0, 10))
        tk.Radiobutton(selector_frame, text="Gmail ID", variable=self.google_id_type,
                       value="Gmail ID", bg="white").pack(side="left")
        tk.Radiobutton(selector_frame, text="GAID", variable=self.google_id_type,
                       value="GAID", bg="white").pack(side="left")
        self.google_id_type.trace_add("write", lambda *args: self.renumber_google_labels())

    def renumber_google_labels(self):
        if self.platform_type == "Google":
            for i, fr in enumerate(self.entry_frames):
                fr.winfo_children()[0].config(text=f"{self.google_id_type.get()} {i + 1}:")

    def on_platform_change(self, event=None):
        platform = self.inter_option.get()
        self.platform_type = platform
        self.entries.clear()

        for fr in self.entry_frames:
            fr.destroy()
        self.entry_frames.clear()

        for widget in self.field_frame.winfo_children():
            widget.destroy()

        if platform == "Google":
            self.show_google_id_selector()

        # ✅ Only add the first field if a valid platform is chosen
        if platform not in (None, "", "Select Platform"):
            self.add_field()


    def update_generate_button_state(self, event=None):
        has_case = self.app.crime_number and self.app.ncrp_id
        has_input = any(e.get().strip() for e in self.entries)
        platform = self.platform_type
        self.inter_generate_button.config(
            state="normal" if has_case and has_input and platform and platform != "Select Platform" else "disabled"
        )

    # ---------------------------
    # (The rest of your generate_inter_letter, generate_inter_word_letter, view_letters_inter remain as before)
    # ---------------------------

    # keep your generate_inter_letter, generate_inter_word_letter, view_letters_inter from before
    def build_accounts_table(self, doc, platform, accounts):
        platform_heading_map = {
            "WhatsApp": "WhatsApp Accounts",
            "Facebook": "URLs",
            "Instagram": "Instagram Accounts",
            "Twitter": "Twitter Accounts",
            "Google": f"{self.google_id_type.get()}s"
        }
        col_title = platform_heading_map.get(platform, "Accounts")

        left_block = accounts[:9]
        right_block = accounts[9:]

        cols = 2 if not right_block else 4
        rows = 1 + max(len(left_block), len(right_block))

        table = doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        # Header
        table.cell(0, 0).text = "S.No"
        table.cell(0, 1).text = col_title
        if cols == 4:
            table.cell(0, 2).text = "S.No"
            table.cell(0, 3).text = col_title

        # Fill left block
        for i, val in enumerate(left_block, start=1):
            table.cell(i, 0).text = str(i)
            table.cell(i, 1).text = val

        # Fill right block
        if cols == 4:
            for j, val in enumerate(right_block, start=10):
                table.cell(j - 9, 2).text = str(j)
                table.cell(j - 9, 3).text = val

        return table   # ✅ IMPORTANT — return the table object


    def generate_inter_letter(self):
        logging.debug("Starting intermediary letter generation")
        if not self.app.crime_number or not self.app.ncrp_id:
            self.inter_status_label.config(text="Please enter the case details first.", fg=self.app.error_color)
            logging.error("Intermediary letter generation failed: Missing crime number or NCRP ID")
            return

        if self.inter_option.get() == "Select Platform":
            self.inter_status_label.config(text="Please select a platform.", fg=self.app.error_color)
            logging.error("Intermediary letter generation failed: No platform selected")
            return

        # Use the latest template_dir from main_app (no re-prompt needed)
        if not self.app.template_dir:
            self.inter_status_label.config(text="No valid template directory selected.", fg=self.app.error_color)
            logging.error("No template directory selected")
            messagebox.showerror("Error", "Please select a valid template directory using 'Template Directory' in the header.")
            return

        # Update inter_template_dir to match main_app's template_dir
        self.inter_template_dir = os.path.join(self.app.template_dir, 'inter')

        platform = self.inter_option.get()
        from_date = self.inter_from_date_entry.get().strip()
        to_date = self.inter_to_date_entry.get().strip()

        if not TKCALENDAR_AVAILABLE:
            if from_date and not re.match(r'^\d{2}-\d{2}-\d{4}$', from_date):
                self.inter_status_label.config(text="Invalid From Date format. Use DD-MM-YYYY.", fg=self.app.error_color)
                logging.error("Invalid From Date format")
                return
            if to_date and not re.match(r'^\d{2}-\d{2}-\d{4}$', to_date):
                self.inter_status_label.config(text="Invalid To Date format. Use DD-MM-YYYY.", fg=self.app.error_color)
                logging.error("Invalid To Date format")
                return

        if platform == "Google":
            if not any(entry.get().strip() for entry in self.entries):
                self.inter_status_label.config(text=f"Please enter at least one {self.google_id_type.get()}.", fg=self.app.error_color)
                logging.error(f"Intermediary letter generation failed: No {self.google_id_type.get()} provided")
                return
        else:
            if not any(entry.get().strip() for entry in self.entries):
                self.inter_status_label.config(text="Please enter at least one URL.", fg=self.app.error_color)
                logging.error("Intermediary letter generation failed: No URLs provided")
                return


        def convert_to_yyyy_mm_dd(date_str):
            if date_str and date_str != 'N/A':
                try:
                    return datetime.strptime(date_str, "%d-%m-%Y").strftime("%Y-%m-%d")
                except ValueError:
                    return 'Invalid Date'
            return 'N/A'

        case = {
            'CrimeNumber': self.app.crime_number,
            'NCRP_ID': self.app.ncrp_id,
            'RecipientName': 'N/A',
            'RequestDate': datetime.now().strftime("%d-%m-%Y"),
            'Platform': platform or 'N/A',
            'AccountID': [entry.get() for entry in self.entries if entry.get()],
            'Address': 'N/A',
            'Date_From': convert_to_yyyy_mm_dd(from_date) or 'N/A',  # Convert if needed
            'Date_To': convert_to_yyyy_mm_dd(to_date) or 'N/A',
            'LetterType': 'Intermediary'
        }

        self.app.date_from = from_date if from_date != 'N/A' else None
        self.app.date_to = to_date if to_date != 'N/A' else None

        save_error = save_case({'CrimeNumber': self.app.crime_number, 'NCRP_ID': self.app.ncrp_id}, self.app.officer['Id'], 'Intermediary')
        if save_error:
            self.inter_status_label.config(text=f"Database error: {save_error}", fg=self.app.error_color)
            messagebox.showerror("Error", f"Failed to save case: {save_error}")
            logging.error(f"Database error: {save_error}")
            return

        output_dir = os.path.join(Path.home(), 'Documents', 'GeneratedLetters', 'inter')
        output_path = os.path.join(output_dir, f"Notice_{case['Platform'].replace(' ', '_')}.docx")

        try:
            self.generate_inter_word_letter(case, output_path)
            messagebox.showinfo("Success", f"Generated letter at {output_path}")
            self.inter_status_label.config(text="Letter generated successfully", fg=self.app.success_color)
            self.view_letters_inter_button.config(state="normal")
            logging.debug(f"Generated intermediary letter: {output_path}")
        except Exception as e:
            self.inter_status_label.config(text=f"Error: {str(e)}", fg=self.app.error_color)
            messagebox.showerror("Error", f"Failed to generate letter: {str(e)}")
            logging.error(f"Failed to generate intermediary letter: {str(e)}")

    def generate_inter_word_letter(self, case, output_path):
        logging.debug(f"Generating intermediary word letter: {output_path}")
        self.app.fetch_officer_details()

        def ensure_dd_mm_yyyy(date_str):
            if date_str and date_str != 'N/A':
                try:
                    return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d-%m-%Y")
                except ValueError:
                    return date_str
            return 'N/A'

        # ✅ Removed '{{Platform_Account_Table}}' from replacements
        replacements = {
            '{{Officer_Name}}': self.app.officer.get('OfficerName', 'Unknown Officer'),
            '{{Officer_Designation}}': self.app.officer.get('Designation', 'Unknown Designation'),
            '{{Officer_Phone}}': self.app.officer.get('Phone', 'N/A'),
            '{{Officer_Email}}': self.app.officer.get('Email', 'N/A'),
            '{{Letter_Date}}': datetime.now().strftime("%d-%m-%Y"),
            '{{Nodal_Officer}}': case.get('RecipientName', 'N/A'),
            '{{Platform_Name}}': case.get('Platform', 'N/A'),
            '{{Platform_Email}}': 'N/A',
            '{{Crime_No_with_Section}}': case.get('CrimeNumber', 'N/A'),
            '{{NCRP_ID}}': case.get('NCRP_ID', 'N/A'),
            '{{Date_From}}': ensure_dd_mm_yyyy(case.get('Date_From', 'N/A')),
            '{{Date_To}}': ensure_dd_mm_yyyy(case.get('Date_To', 'N/A'))
        }

        logging.debug(f"Replacements: {replacements}")

        template_map = {
            "Instagram": "instagram_template.docx",
            "WhatsApp": "whatsapp_template.docx",
            "Facebook": "facebook_template.docx",
            "Telegram": "telegram_template.docx",
            "Google": "google_template.docx",
            "Twitter": "twitter_template.docx"
        }

        template_filename = template_map.get(case.get("Platform"), "inter_template.docx")
        template_path = os.path.join(self.inter_template_dir, template_filename)

        if not os.path.exists(template_path) and getattr(sys, 'frozen', False):
            template_path = os.path.join(sys._MEIPASS, 'templates', 'inter', template_filename)

        if not os.path.exists(template_path):
            self.inter_status_label.config(
                text=f"Template file '{template_filename}' not found in {self.inter_template_dir}",
                fg=self.app.error_color
            )
            logging.error(f"Template file not found: {template_path}")
            messagebox.showerror("Error", f"Template file '{template_filename}' not found.")
            return

        try:
            doc = Document(template_path)
        except Exception as e:
            self.inter_status_label.config(
                text=f"Failed to load template: {str(e)}",
                fg=self.app.error_color
            )
            messagebox.showerror("Error", f"Failed to load template '{template_path}': {str(e)}")
            logging.error(f"Failed to load template: {str(e)}")
            raise

        # 🚀 TABLE INSERTION BLOCK — placed right after loading the doc and before other replacements
        placeholder = None
        for para in doc.paragraphs:
            if "{{Platform_Account_Table}}" in para.text:
                placeholder = para
                break

        accounts = case.get('AccountID', [])
        if accounts and placeholder is not None:
            # Remove just the placeholder from the paragraph
            placeholder.text = placeholder.text.replace("{{Platform_Account_Table}}", "").strip()

            # Build the actual table
            tbl = self.build_accounts_table(doc, case.get("Platform", "N/A"), accounts)

            # Insert table XML after this paragraph
            placeholder._p.addnext(tbl._tbl)

        elif accounts and placeholder is None:
            logging.warning("Table placeholder '{{Platform_Account_Table}}' not found in template.")

        # 🔄 Now perform the other placeholder replacements
        for paragraph in doc.paragraphs:
            replace_placeholder_in_paragraph(paragraph, replacements)

        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, replacements)

        # 💾 Save the final document
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            logging.debug(f"Saved intermediary letter: {output_path}")
        except Exception as e:
            self.inter_status_label.config(
                text=f"Failed to save letter: {str(e)}",
                fg=self.app.error_color
            )
            messagebox.showerror("Error", f"Failed to save letter to '{output_path}': {str(e)}")
            logging.error(f"Failed to save letter: {str(e)}")
            raise


    def view_letters_inter(self):
        logging.debug("Opening intermediary letters folder")
        folder_path = os.path.join(Path.home(), 'Documents', 'GeneratedLetters', 'inter')
        try:
            if os.path.exists(folder_path):
                os.startfile(folder_path)
                logging.debug(f"Opened intermediary letters folder: {folder_path}")
            else:
                messagebox.showwarning("Warning", "No letters generated or folder not found")
                logging.warning(f"Intermediary letters folder not found: {folder_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open folder: {str(e)}")
            logging.error(f"Failed to open intermediary letters folder: {str(e)}")
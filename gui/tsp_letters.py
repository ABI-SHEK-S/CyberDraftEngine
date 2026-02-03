import logging
import os
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from datetime import datetime
import re
from docx import Document
import sys

try:
    from tkcalendar import DateEntry
    TKCALENDAR_AVAILABLE = True
except ImportError:
    TKCALENDAR_AVAILABLE = False
from .utils import replace_placeholder_in_paragraph
from db.database import save_case


class TSPLetters:
    def __init__(self, parent, app):
        self.parent = parent
        self.app = app
        self.input_entries = []
        # Initialize logging
        log_dir = os.path.join(Path.home(), 'Documents', 'LetterGeneratorLogs')
        os.makedirs(log_dir, exist_ok=True)
        logging.basicConfig(
            filename=os.path.join(log_dir, 'app.log'),
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        logging.debug("TSPLetters initialized")
        # Use app's template directory
        self.tsp_template_dir = os.path.join(self.app.template_dir, 'tsp') if self.app.template_dir else None
        self.setup_ui()

    def setup_ui(self):
        logging.debug("Setting up TSP Letters UI")
        tsp_inner = tk.Frame(self.parent, bg="white", bd=2, relief="flat")
        tsp_inner.pack(pady=20, padx=60, fill="both", expand=True)
        tsp_inner.grid_columnconfigure(0, weight=1)
        tsp_inner.grid_columnconfigure(1, weight=1)
        tsp_inner.grid_rowconfigure(2, weight=1)

        tk.Label(tsp_inner, text="TSP:", bg="white", font=("Segoe UI", 10, "bold"), fg=self.app.text_color).grid(row=0, column=0, sticky="w", pady=5)
        self.tsp_option = ttk.Combobox(
            tsp_inner, values=["Airtel", "Jio", "Vodafone", "BSNL"], state="readonly", style="TCombobox"
        )
        self.tsp_option.set("Select TSP")
        self.tsp_option.grid(row=1, column=0, sticky="ew", pady=5)
        self.app.ToolTip(self.tsp_option, "Select the Telecom Service Provider for the letter", self.app)

        tk.Label(tsp_inner, text="Request Type:", bg="white", font=("Segoe UI", 10, "bold"), fg=self.app.text_color).grid(row=0, column=1, sticky="w", pady=5)
        self.request_type_option = ttk.Combobox(
            tsp_inner, values=["CAF", "CDR", "IMEI CDR", "Aadhar linked numbers", "PoS code"], state="readonly", style="TCombobox"
        )
        self.request_type_option.set("Select Request Type")
        self.request_type_option.grid(row=1, column=1, sticky="ew", pady=5, padx=(5, 0))
        self.request_type_option.bind("<<ComboboxSelected>>", self.toggle_input_fields)
        self.app.ToolTip(self.request_type_option, "Select the request type for the TSP letter", self.app)

        # Canvas with scrollbar
        self.input_canvas = tk.Canvas(tsp_inner, bg="white", highlightthickness=0)
        self.input_canvas.grid(row=2, column=0, sticky="nsew", pady=5, columnspan=2)
        scrollbar = ttk.Scrollbar(tsp_inner, orient="vertical", command=self.input_canvas.yview)
        scrollbar.grid(row=2, column=2, sticky="ns")
        self.input_canvas.configure(yscrollcommand=scrollbar.set)

        self.input_entry_frame = tk.Frame(self.input_canvas, bg="white")
        self._input_window_id = self.input_canvas.create_window((0, 0), window=self.input_entry_frame, anchor="nw")

        self.input_entry_frame.bind("<Configure>", lambda e: self.input_canvas.configure(scrollregion=self.input_canvas.bbox("all")))
        self.input_canvas.bind("<Configure>", self._resize_input_frame)

        # Mouse wheel scroll binding
        self._bind_mousewheel(self.input_entry_frame)
        self._bind_mousewheel(self.input_canvas)

        self.input_entries = []

        self.tsp_date_frame = tk.Frame(tsp_inner, bg="white")
        self.tsp_date_frame.grid(row=3, column=0, sticky="ew", pady=5, columnspan=2)
        tk.Label(self.tsp_date_frame, text="Date Range:", font=("Segoe UI", 10, "bold"), bg="white", fg=self.app.text_color).pack(anchor="w", pady=(5, 2))
        date_inner = tk.Frame(self.tsp_date_frame, bg="white")
        date_inner.pack(anchor="w")
        self.tsp_from_date_label = tk.Label(date_inner, text="From Date:", bg="white", fg=self.app.text_color)
        self.tsp_from_date_label.pack(side="left", padx=5)
        if TKCALENDAR_AVAILABLE:
            self.tsp_from_date_entry = DateEntry(date_inner, date_pattern='dd-mm-yyyy', width=12, background=self.app.accent_color, foreground='white')
            self.app.ToolTip(self.tsp_from_date_entry, "Select the start date (DD-MM-YYYY) for the TSP request", self.app)
        else:
            self.tsp_from_date_entry = ttk.Entry(date_inner, width=12, style="TEntry")
            self.app.ToolTip(self.tsp_from_date_entry, "Enter date in DD-MM-YYYY format (tkcalendar not installed)", self.app)
        self.tsp_from_date_entry.pack(side="left", padx=5)
        self.tsp_to_date_label = tk.Label(date_inner, text="To Date:", bg="white", fg=self.app.text_color)
        self.tsp_to_date_label.pack(side="left", padx=5)
        if TKCALENDAR_AVAILABLE:
            self.tsp_to_date_entry = DateEntry(date_inner, date_pattern='dd-mm-yyyy', width=12, background=self.app.accent_color, foreground='white')
            self.app.ToolTip(self.tsp_to_date_entry, "Select the end date (DD-MM-YYYY) for the TSP request", self.app)
        else:
            self.tsp_to_date_entry = ttk.Entry(date_inner, width=12, style="TEntry")
            self.app.ToolTip(self.tsp_to_date_entry, "Enter date in DD-MM-YYYY format (tkcalendar not installed)", self.app)
        self.tsp_to_date_entry.pack(side="left", padx=5)
        self.tsp_date_frame.grid_forget()

        button_frame = tk.Frame(tsp_inner, bg="white")
        button_frame.grid(row=4, column=0, columnspan=2, pady=10)

        self.tsp_generate_button = ttk.Button(
            button_frame, text="Generate Letter", command=self.generate_tsp_letter,
            style="TButton", state="disabled", width=20
        )
        self.tsp_generate_button.pack(pady=5)
        self.app.ToolTip(self.tsp_generate_button, "Generate the TSP letter", self.app)

        self.view_letters_tsp_button = ttk.Button(
            button_frame, text="View Letters", command=self.view_letters_tsp,
            style="TButton", state="disabled", width=20
        )
        self.view_letters_tsp_button.pack(pady=5)
        self.app.ToolTip(self.view_letters_tsp_button, "Open the folder containing generated TSP letters", self.app)

        self.tsp_status_label = tk.Label(
            tsp_inner, text="", font=("Segoe UI", 10), bg="white", fg=self.app.success_color
        )
        self.tsp_status_label.grid(row=5, column=0, sticky="w", pady=5, columnspan=2)

        logging.debug("TSP Letters UI setup complete")

    def _bind_mousewheel(self, widget):
        widget.bind_all("<MouseWheel>", self._on_mousewheel, add="+")
        widget.bind_all("<Button-4>", self._on_mousewheel, add="+")
        widget.bind_all("<Button-5>", self._on_mousewheel, add="+")

    def _on_mousewheel(self, event):
        if event.num == 4:
            self.input_canvas.yview_scroll(-1, "units")
        elif event.num == 5:
            self.input_canvas.yview_scroll(1, "units")
        else:
            delta = -1 if event.delta > 0 else 1
            self.input_canvas.yview_scroll(delta, "units")

    def _resize_input_frame(self, event):
        canvas_width = event.width
        self.input_canvas.itemconfigure(self._input_window_id, width=canvas_width)
        logging.debug("Resized input canvas")

    def _scroll_to_widget(self, widget):
        try:
            self.input_canvas.update_idletasks()
            y = widget.winfo_y()
            height = widget.winfo_height()
            frame_height = self.input_entry_frame.winfo_height()
            target = max(0.0, min(1.0, (y + height) / max(1, frame_height)))
            self.input_canvas.yview_moveto(target)
        except Exception as e:
            logging.warning(f"Auto-scroll failed: {e}")

    def toggle_input_fields(self, event=None):
        for entry in self.input_entries:
            entry.destroy()
        self.input_entries.clear()
        self.tsp_date_frame.grid_forget()
        request_type = self.request_type_option.get()
        if request_type == "Select Request Type":
            return
        if request_type in ["CAF", "CDR", "IMEI CDR", "Aadhar linked numbers", "PoS code"]:
            self.add_input_field(request_type)
            if request_type in ["CDR", "IMEI CDR"]:
                self.tsp_date_frame.grid(row=3, column=0, pady=5, sticky="nsew", columnspan=2)
        self.input_canvas.yview_moveto(0.0)
        logging.debug(f"Toggled input fields for request type: {request_type}")

    def add_input_field(self, request_type):
        frame = tk.Frame(self.input_entry_frame, bg="white")
        frame.pack(pady=5, fill="x")

        number = len(self.input_entries) + 1
        if request_type in ["CAF", "CDR"]:
            label = tk.Label(frame, text=f"Phone No {number}:", bg="white", fg=self.app.text_color)
            entry = ttk.Entry(frame, width=20, style="TEntry")
            self.app.ToolTip(entry, "Enter the phone number (10 digits)", self.app)
        elif request_type == "IMEI CDR":
            label = tk.Label(frame, text=f"IMEI No {number}:", bg="white", fg=self.app.text_color)
            entry = ttk.Entry(frame, width=20, style="TEntry")
            self.app.ToolTip(entry, "Enter the IMEI number (15 digits)", self.app)
        elif request_type == "Aadhar linked numbers":
            label = tk.Label(frame, text=f"Aadhar No {number}:", bg="white", fg=self.app.text_color)
            entry = ttk.Entry(frame, width=20, style="TEntry")
            self.app.ToolTip(entry, "Enter the Aadhar number (12 digits)", self.app)
        elif request_type == "PoS code":
            label = tk.Label(frame, text=f"PoS Code {number}:", bg="white", fg=self.app.text_color)
            entry = ttk.Entry(frame, width=20, style="TEntry")
            self.app.ToolTip(entry, "Enter the PoS code", self.app)

        label.pack(side="left", padx=5)
        entry.pack(side="left", padx=5)

        add_button = ttk.Button(frame, text="+", command=lambda: self.add_input_field(request_type), style="TButton")
        add_button.pack(side="left", padx=2)
        self.app.ToolTip(add_button, "Add another input field", self.app)

        delete_button = ttk.Button(frame, text="Delete", command=lambda: self.delete_input_field(frame), style="TButton")
        delete_button.pack(side="left", padx=2)
        self.app.ToolTip(delete_button, "Delete this input field", self.app)

        if len(self.input_entries) == 0:
            delete_button.config(state="disabled")

        self.input_entries.append(frame)
        self.input_canvas.configure(scrollregion=self.input_canvas.bbox("all"))
        logging.debug(f"Added input field for request type: {request_type}")

        self.tsp_generate_button.config(state="normal" if self.app.crime_number and self.app.ncrp_id else "disabled")

        self.update_input_numbering(request_type)

        self._scroll_to_widget(frame)

    def delete_input_field(self, frame):
        if len(self.input_entries) > 1:
            if frame in self.input_entries:
                self.input_entries.remove(frame)
            frame.destroy()
            if len(self.input_entries) == 1:
                self.input_entries[0].winfo_children()[-1].config(state="disabled")
            self.input_canvas.configure(scrollregion=self.input_canvas.bbox("all"))
            logging.debug("Deleted input field")
            if self.request_type_option.get() != "Select Request Type":
                self.update_input_numbering(self.request_type_option.get())
        if not self.input_entries:
            self.tsp_generate_button.config(state="disabled")

    def update_input_numbering(self, request_type):
        for idx, frame in enumerate(self.input_entries, start=1):
            label = frame.winfo_children()[0]
            base_label = ""
            if request_type in ["CAF", "CDR"]:
                base_label = "Phone No"
            elif request_type == "IMEI CDR":
                base_label = "IMEI No"
            elif request_type == "Aadhar linked numbers":
                base_label = "Aadhar No"
            elif request_type == "PoS code":
                base_label = "PoS Code"
            label.config(text=f"{base_label} {idx}:")

    # --- rest of your original methods (generate_tsp_letter, generate_tsp_word_letter, view_letters_tsp) remain unchanged ---

    def generate_tsp_letter(self):
        logging.debug("Starting TSP letter generation")
        if not self.app.crime_number or not self.app.ncrp_id:
            self.tsp_status_label.config(text="Please enter the case details first.", fg=self.app.error_color)
            logging.error("TSP letter generation failed: Missing crime number or NCRP ID")
            return

        # Use the latest template_dir from main_app (no re-prompt needed)
        if not self.app.template_dir:
            self.tsp_status_label.config(text="No valid template directory selected.", fg=self.app.error_color)
            logging.error("No template directory selected")
            messagebox.showerror("Error", "Please select a valid template directory using 'Template Directory' in the header.")
            return
        

        # Update tsp_template_dir to match main_app's template_dir
        self.tsp_template_dir = os.path.join(self.app.template_dir, 'tsp')

        if not self.tsp_option.get() or self.tsp_option.get() == "Select TSP":
            self.tsp_status_label.config(text="Please select a TSP.", fg=self.app.error_color)
            logging.error("TSP letter generation failed: No TSP selected")
            return

        request_type = self.request_type_option.get()
        if request_type == "Select Request Type":
            self.tsp_status_label.config(text="Please select a request type.", fg=self.app.error_color)
            logging.error("TSP letter generation failed: No request type selected")
            return

        inputs = [entry.winfo_children()[1].get().strip() for entry in self.input_entries]
        if not any(inputs):
            self.tsp_status_label.config(text=f"At least one {request_type.lower()} is required.", fg=self.app.error_color)
            logging.error(f"TSP letter generation failed: No {request_type.lower()} provided")
            return

        if request_type in ["CAF", "CDR"]:
            if not all(re.match(r'^\d{10}$', i) for i in inputs if i):
                self.tsp_status_label.config(text="Phone No must be 10 digits.", fg=self.app.error_color)
                logging.error("TSP letter generation failed: Invalid phone number format")
                return

        if request_type == "IMEI CDR":
            if not all(re.match(r'^\d{15}$', i) for i in inputs if i):
                self.tsp_status_label.config(text="IMEI No must be 15 digits.", fg=self.app.error_color)
                logging.error("TSP letter generation failed: Invalid IMEI number format")
                return

        if request_type == "Aadhar linked numbers":
            if not all(re.match(r'^\d{12}$', i) for i in inputs if i):
                self.tsp_status_label.config(text="Aadhar No must be 12 digits.", fg=self.app.error_color)
                logging.error("TSP letter generation failed: Invalid Aadhar number format")
                return

        if request_type == "PoS code":
            if not all(i for i in inputs if i):
                self.tsp_status_label.config(text="PoS Code cannot be empty.", fg=self.app.error_color)
                logging.error("TSP letter generation failed: Empty PoS code")
                return

        from_date = self.tsp_from_date_entry.get().strip()
        to_date = self.tsp_to_date_entry.get().strip()

        if request_type in ["CDR", "IMEI CDR"]:
            if not from_date or not to_date:
                self.tsp_status_label.config(text="From and To dates are required.", fg=self.app.error_color)
                logging.error("TSP letter generation failed: Missing date range")
                return
            if not TKCALENDAR_AVAILABLE:
                if not re.match(r'^\d{2}-\d{2}-\d{4}$', from_date) or not re.match(r'^\d{2}-\d{2}-\d{4}$', to_date):
                    self.tsp_status_label.config(text="Dates must be in DD-MM-YYYY format.", fg=self.app.error_color)
                    logging.error("TSP letter generation failed: Invalid date format")
                    return

        case = {
            'CrimeNumber': self.app.crime_number,
            'NCRP_ID': self.app.ncrp_id,
            'RecipientName': 'N/A',
            'RequestDate': datetime.now().strftime("%d-%m-%Y"),
            'TSP': self.tsp_option.get() or 'N/A',
            'MobileNo': inputs if request_type in ["CAF", "CDR"] else [],
            'Address': 'N/A',
            'Date_Ranges': [(from_date, to_date)] if request_type in ["CDR", "IMEI CDR"] else [],
            'LetterType': 'TSP',
            'Request_Type': request_type,
            'IMEI_No': inputs if request_type == "IMEI CDR" else [],
            'Aadhar_No': inputs if request_type == "Aadhar linked numbers" else [],
            'PoS_Code': inputs if request_type == "PoS code" else []
        }

        logging.debug(f"Case dictionary: {case}")

        self.app.date_from = from_date if from_date != 'N/A' else None
        self.app.date_to = to_date if to_date != 'N/A' else None

        save_error = save_case({'CrimeNumber': self.app.crime_number, 'NCRP_ID': self.app.ncrp_id}, self.app.officer.get('Id', 'N/A'), 'TSP')
        if save_error:
            self.tsp_status_label.config(text=f"Database error: {save_error}", fg=self.app.error_color)
            logging.error(f"Database error: {save_error}")
            return

        output_dir = os.path.join(Path.home(), 'Documents', 'GeneratedLetters', 'tsp')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"Notice_{case['TSP'].replace(' ', '_')}_{case['Request_Type'].replace(' ', '_')}.docx")

        try:
            self.generate_tsp_word_letter(case, output_path)
            messagebox.showinfo("Success", f"Generated letter at {output_path}")
            self.tsp_status_label.config(text="Letter generated successfully", fg=self.app.success_color)
            self.view_letters_tsp_button.config(state="normal")
            logging.debug(f"TSP letter generated: {output_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate letter: {str(e)}")
            self.tsp_status_label.config(text=f"Error: {str(e)}", fg=self.app.error_color)
            logging.error(f"Failed to generate TSP letter: {str(e)}")

    def generate_tsp_word_letter(self, case, output_path):
        logging.debug(f"Generating TSP word letter: {output_path}")

        if not hasattr(self.app, 'fetch_officer_details'):
            self.tsp_status_label.config(text="Application error: Officer details unavailable.", fg=self.app.error_color)
            messagebox.showerror("Error", "Application error: Officer details unavailable.")
            logging.error("Officer details unavailable in app")
            raise AttributeError("fetch_officer_details not found in app")

        self.app.fetch_officer_details()

        input_list = (
            case.get('MobileNo', ['N/A']) if case.get('Request_Type') in ["CAF", "CDR"] else
            case.get('IMEI_No', ['N/A']) if case.get('Request_Type') == "IMEI CDR" else
            case.get('Aadhar_No', ['N/A']) if case.get('Request_Type') == "Aadhar linked numbers" else
            case.get('PoS_Code', ['N/A']) if case.get('Request_Type') == "PoS code" else ['N/A']
        )

        replacements = {
            '{{Officer_Name}}': self.app.officer.get('OfficerName', 'Unknown Officer'),
            '{{Officer_Designation}}': self.app.officer.get('Designation', 'Unknown Designation'),
            '{{Officer_Phone}}': self.app.officer.get('Phone', 'N/A'),
            '{{Officer_Email}}': self.app.officer.get('Email', 'N/A'),
            '{{Letter_Date}}': datetime.now().strftime("%d-%m-%Y"),
            '{{Nodal_Officer}}': case.get('RecipientName', 'N/A'),
            '{{Platform_Name}}': case.get('TSP', 'N/A'),
            '{{Platform_Email}}': 'N/A',
            '{{Crime_No_with_Section}}': case.get('CrimeNumber', 'N/A'),
            '{{NCRP_ID}}': case.get('NCRP_ID', 'N/A'),
            '{{Platform_Account_Table}}': ', '.join(input_list),
            '{{Date_From}}': case.get('Date_Ranges', [('N/A', 'N/A')])[0][0] if case.get('Date_Ranges') else 'N/A',
            '{{Date_To}}': case.get('Date_Ranges', [('N/A', 'N/A')])[0][1] if case.get('Date_Ranges') else 'N/A',
            '{{Request_Type}}': case.get('Request_Type', 'N/A')
        }

        logging.debug(f"Replacements: {replacements}")

        template_map = {
            "CAF": "caf_template.docx",
            "CDR": "cdr_template.docx",
            "IMEI CDR": "imei_cdr_template.docx",
            "Aadhar linked numbers": "aadhar_template.docx",
            "PoS code": "pos_template.docx"
        }

        template_filename = template_map.get(case.get("Request_Type"), None)
        if not template_filename:
            self.tsp_status_label.config(text=f"Invalid request type: {case.get('Request_Type')}", fg=self.app.error_color)
            messagebox.showerror("Error", f"Invalid request type: {case.get('Request_Type')}")
            logging.error(f"Invalid request type: {case.get('Request_Type')}")
            return
        # Build path using selected dir first
        template_path = os.path.join(self.tsp_template_dir, template_filename)

        # Fallback to EXE temp path only if selected path doesn't exist
        if not os.path.exists(template_path) and getattr(sys, 'frozen', False):
            template_path = os.path.join(sys._MEIPASS, 'templates', 'tsp', template_filename)

        # Check if the file exists
        if not os.path.exists(template_path):
            self.tsp_status_label.config(text=f"Template file '{template_filename}' not found in {self.tsp_template_dir}", fg=self.app.error_color)
            logging.error(f"Template file not found: {template_path}")
            messagebox.showerror("Error", f"Template file '{template_filename}' not found in {self.tsp_template_dir}. Ensure your selected directory contains 'tsp/{template_filename}'.")
            return

        try:
            doc = Document(template_path)
        except Exception as e:
            self.tsp_status_label.config(text=f"Failed to load template: {str(e)}", fg=self.app.error_color)
            messagebox.showerror("Error", f"Failed to load template '{template_path}': {str(e)}")
            logging.error(f"Failed to load template: {str(e)}")
            raise

        for paragraph in doc.paragraphs:
            replace_placeholder_in_paragraph(paragraph, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, replacements)

        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            logging.debug(f"Saved TSP letter: {output_path}")
        except Exception as e:
            self.tsp_status_label.config(text=f"Failed to save letter: {str(e)}", fg=self.app.error_color)
            messagebox.showerror("Error", f"Failed to save letter to '{output_path}': {str(e)}")
            logging.error(f"Failed to save letter: {str(e)}")
            raise


    def view_letters_tsp(self):
        logging.debug("Opening TSP letters folder")
        folder_path = os.path.join(Path.home(), 'Documents', 'GeneratedLetters', 'tsp')
        try:
            if os.path.exists(folder_path):
                os.startfile(folder_path)
                logging.debug(f"Opened TSP letters folder: {folder_path}")
            else:
                messagebox.showwarning("Warning", "No letters generated or folder not found")
                logging.warning(f"TSP letters folder not found: {folder_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open folder: {str(e)}")
            logging.error(f"Failed to open TSP letters folder: {str(e)}")
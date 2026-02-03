from pathlib import Path
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
import os
from docx import Document
from db.database import connect_db, save_case, validate_case
from datetime import datetime
from .utils import replace_placeholder_in_paragraph
import re
import logging

class BankLetters:
    def __init__(self, parent, app):
        self.parent = parent
        self.app = app
        log_dir = os.path.join(Path.home(), 'Documents', 'LetterGeneratorLogs')
        os.makedirs(log_dir, exist_ok=True)
        logging.basicConfig(
            filename=os.path.join(log_dir, 'bank_letters.log'),
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        logging.debug("BankLetters initialized")
        self.bank_template_dir = None   # will be set each time a letter is generated
        self.setup_ui()

    def setup_ui(self):
        bank_inner = tk.Frame(self.parent, bg="white", bd=2, relief="flat")
        bank_inner.pack(pady=20, padx=60, fill="x")

        self.excel_button = ttk.Button(bank_inner, text="Browse Excel File", command=self.select_excel, style="TButton")
        self.excel_button.pack(pady=5)
        self.app.ToolTip(self.excel_button, "Select an Excel file to process bank letter data", self.app)

        self.excel_label = tk.Label(bank_inner, text="No file selected", font=("Segoe UI", 10), bg="white", fg=self.app.text_color)
        self.excel_label.pack(pady=5)

        self.progress_bar = ttk.Progressbar(bank_inner, mode='determinate', length=300)
        self.progress_bar.pack(pady=5, fill="x")
        self.progress_bar.pack_forget()

        self.generate_button = ttk.Button(
            bank_inner, text="Generate Letters ", command=self.process_excel,
            style="TButton", state="disabled"
        )
        self.generate_button.pack(pady=5)
        self.app.ToolTip(self.generate_button, "Generate letters from the selected Excel file ", self.app)

        self.view_letters_bank_button = ttk.Button(
            bank_inner, text="View Letters", command=self.view_letters_bank,
            style="TButton", state="disabled"
        )
        self.view_letters_bank_button.pack(pady=5)
        self.app.ToolTip(self.view_letters_bank_button, "Open the folder containing generated bank letters", self.app)

        self.bank_status_label = tk.Label(
            bank_inner, text="", font=("Segoe UI", 10), bg="white", fg=self.app.success_color
        )
        self.bank_status_label.pack(pady=5)
        logging.debug("BankLetters UI setup complete")

    def format_inr(self, amount):
        try:
            amount = float(re.sub(r'[^\d.]', '', str(amount)))
        except:
            return str(amount)
        s = f"{amount:,.2f}"
        s = s.split('.')
        x = s[0]
        if len(x) > 3:
            x = x[:-3].replace(',', '')[::-1]
            x = ','.join([x[i:i+2] for i in range(0, len(x), 2)])[::-1] + ',' + s[0][-3:]
        else:
            x = s[0]
        return f"₹{x}/-"

    def clean_account_number(self, x):
        try:
            if isinstance(x, float) and x.is_integer():
                return str(int(x))
            if isinstance(x, str) and x.replace('.', '', 1).isdigit():
                f = float(x)
                if f.is_integer():
                    return str(int(f))
                return str(f)
            return str(x)
        except Exception:
            return str(x)

    def select_excel(self):
        self.selected_file = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
        if self.selected_file:
            try:
                df = pd.read_excel(self.selected_file, sheet_name=0)
                df.columns = [col.strip().lower().replace(' ', '_') for col in df.columns]
                required_columns = {
                    'account_no', 'ifsc_code', 'transaction_amount', 'date_from',
                    'date_to', 'transaction_id_/_utr_number2', 'bank/fis'
                }
                missing_columns = required_columns - set(df.columns)
                if missing_columns:
                    self.bank_status_label.config(text=f"Invalid Excel: Missing columns {', '.join(missing_columns)}", fg=self.app.error_color)
                    self.generate_button.config(state="disabled")
                    messagebox.showerror("Error", f"Missing required columns: {', '.join(missing_columns)}")
                    logging.error(f"Missing Excel columns: {missing_columns}")
                    return
                if df.empty:
                    self.bank_status_label.config(text="Invalid Excel: Sheet is empty", fg=self.app.error_color)
                    self.generate_button.config(state="disabled")
                    messagebox.showerror("Error", "Excel sheet is empty")
                    logging.error("Excel sheet is empty")
                    return
                self.excel_label.config(text=f"Selected: {os.path.basename(self.selected_file)}")
                self.app.update_button_states()
                logging.debug(f"Selected Excel file: {self.selected_file}")
            except pd.errors.EmptyDataError:
                self.bank_status_label.config(text="Invalid Excel: File is empty or corrupted", fg=self.app.error_color)
                self.generate_button.config(state="disabled")
                messagebox.showerror("Error", "Excel file is empty or corrupted")
                logging.error("Excel file is empty or corrupted")
            except Exception as e:
                self.bank_status_label.config(text=f"Invalid Excel file: {str(e)}", fg=self.app.error_color)
                self.generate_button.config(state="disabled")
                messagebox.showerror("Error", f"Failed to read Excel file: {str(e)}")
                logging.error(f"Failed to read Excel file: {str(e)}")
        else:
            self.excel_label.config(text="No file selected")
            self.generate_button.config(state="disabled")
            logging.debug("No Excel file selected")

    def process_excel(self):
        logging.debug("Starting bank letter generation")
        if not self.app.crime_number or not self.app.ncrp_id:
            self.bank_status_label.config(text="Please enter both crime number and NCRP ID.", fg=self.app.error_color)
            logging.error("Bank letter generation failed: Missing crime number or NCRP ID")
            return
# Note: Removed check for ncrp_id to allow any value (including empty)


        if not self.selected_file:
            self.bank_status_label.config(text="Please select an Excel file", fg=self.app.error_color)
            logging.error("No Excel file selected")
            return

        if not self.app.template_dir:
            self.bank_status_label.config(text="No valid template directory selected.", fg=self.app.error_color)
            logging.error("No template directory selected")
            messagebox.showerror("Error", "Please select a valid template directory using 'Template Directory' in the header.")
            return

        # Always pick up the folder the user selected in the header
        self.bank_template_dir = os.path.join(self.app.template_dir, "banks")

        # 1) Try the user-selected folder first
        template_path = os.path.join(self.bank_template_dir, "bank.docx")

        # 2) Fallback for the PyInstaller EXE – only if the first path is missing
        if not os.path.exists(template_path) and getattr(sys, "frozen", False):
            template_path = os.path.join(
                sys._MEIPASS,               # temp folder created by PyInstaller
                "templates", "banks", "bank.docx"
            )

        # 3) Final check – abort with a clear, user-friendly message
        if not os.path.exists(template_path):
            err_msg = (
                "Template file 'bank.docx' not found in:\n"
                f"  {self.bank_template_dir}\n\n"
                "Make sure the folder you chose with “Template Directory” "
                "contains a sub-folder named  banks  and that banks\\bank.docx exists."
            )
            self.bank_status_label.config(text=err_msg, fg=self.app.error_color)
            logging.error(f"Template file not found: {template_path}")
            messagebox.showerror("Template Missing", err_msg)
            return

        # Proceed with loading the document (rest of your code here)
        try:
            df = pd.read_excel(self.selected_file, sheet_name=0)
            if df.empty:
                raise ValueError("The Excel sheet is empty")
            df.columns = [col.strip().lower().replace(' ', '_') for col in df.columns]
            required_columns = {'account_no', 'ifsc_code', 'transaction_amount', 'date_from', 'date_to', 'transaction_id_/_utr_number2', 'bank/fis'}
            missing_columns = required_columns - set(df.columns)
            if missing_columns:
                self.bank_status_label.config(text=f"Missing required columns: {', '.join(missing_columns)}", fg=self.app.error_color)
                messagebox.showerror("Error", f"Missing required columns: {', '.join(missing_columns)}")
                logging.error(f"Missing columns: {missing_columns}")
                return
            total_amount = df['transaction_amount'].apply(lambda x: float(re.sub(r'[^\d.]', '', str(x))) if pd.notnull(x) else 0).sum()
            errors = []
            success_count = 0
            max_letters = 200
            self.progress_bar.pack()
            grouped = df.groupby('bank/fis')
            self.progress_bar['maximum'] = min(len(grouped), max_letters)
            self.progress_bar['value'] = 0
            for bank_name, group in grouped:
                unique_accounts = group[['account_no', 'ifsc_code']].drop_duplicates(subset=['account_no'])
                case = {
                    'CrimeNumber': self.app.crime_number,
                    'NCRP_ID': self.app.ncrp_id,
                    'Total_Amount': self.format_inr(total_amount),
                    'Bank': str(bank_name).strip() or 'Unknown Bank',
                    'RequestDate': datetime.now().strftime("%d-%m-%Y"),
                    'RecipientName': 'Nodal Officer',
                    'Address': 'N/A',
                    'Date_From': group['date_from'].dropna().min().strftime('%d-%m-%Y') if pd.notnull(group['date_from'].dropna().min()) else 'N/A',
                    'Date_To': group['date_to'].dropna().max().strftime('%d-%m-%Y') if pd.notnull(group['date_to'].dropna().max()) else 'N/A',
                    'Accounts': unique_accounts.to_dict('records')
                }
                validation_errors = validate_case({
                    'CrimeNumber': case['CrimeNumber'],
                    'NCRP_ID': case['NCRP_ID'],
                    'AccountNumber': case['Accounts'][0]['account_no'] if case['Accounts'] else 'N/A',
                    'IFSCCode': case['Accounts'][0]['ifsc_code'] if case['Accounts'] else 'N/A',
                    'TransactionAmount': total_amount,
                    'TransactionID': group['transaction_id_/_utr_number2'].iloc[0] if not group['transaction_id_/_utr_number2'].empty else 'N/A',
                    'Bank': case['Bank'],
                    'RequestDate': case['RequestDate'],
                    'RecipientName': case['RecipientName'],
                    'Address': case['Address'],
                    'Date_From': case['Date_From'],
                    'Date_To': case['Date_To']
                })
                if validation_errors:
                    errors.append(f"Bank {bank_name}: Validation warnings - {'; '.join(validation_errors)}")
                    logging.error(f"Validation errors for bank {bank_name}: {validation_errors}")
                    continue
                save_error = save_case({'CrimeNumber': self.app.crime_number, 'NCRP_ID': self.app.ncrp_id}, self.app.officer['Id'], 'Bank')
                if save_error:
                    errors.append(f"Bank {bank_name}: Database error - {save_error}")
                    logging.error(f"Database error for bank {bank_name}: {save_error}")
                    continue
                output_path = os.path.join(Path.home(), 'Documents', 'GeneratedLetters', 'bank', f"Notice_{case['Bank'].replace(' ', '_')}_{success_count + 1}.docx")
                try:
                    self.generate_word_letter(case, output_path)
                    success_count += 1
                    self.progress_bar['value'] = success_count
                    self.app.root.update_idletasks()
                    logging.debug(f"Generated letter for bank {bank_name}: {output_path}")
                except Exception as e:
                    errors.append(f"Bank {bank_name}: Failed to generate letter - {str(e)}")
                    logging.error(f"Failed to generate letter for bank {bank_name}: {str(e)}")
                if success_count >= max_letters:
                    break
            self.progress_bar.pack_forget()
            if success_count > 0:
                messagebox.showinfo("Success", f"Generated {success_count} letters in 'GeneratedLetters/bank' folder")
                self.view_letters_bank_button.config(state="normal")
                self.bank_status_label.config(text=f"Processed {success_count} cases. {len(errors)} issues", fg=self.app.success_color)
                logging.debug(f"Processed {success_count} bank letters with {len(errors)} errors")
            else:
                self.bank_status_label.config(text=f"No letters generated. {len(errors)} issues", fg=self.app.error_color)
                logging.warning(f"No bank letters generated. {len(errors)} errors")
            if errors:
                self.app.show_error_log(errors)
        except FileNotFoundError:
            self.bank_status_label.config(text="Excel file not found", fg=self.app.error_color)
            logging.error("Excel file not found")
            messagebox.showerror("Error", "Excel file not found")
        except ValueError as e:
            self.bank_status_label.config(text=f"Invalid Excel data: {str(e)}", fg=self.app.error_color)
            logging.error(f"Invalid Excel data: {str(e)}")
            messagebox.showerror("Error", f"Invalid Excel data: {str(e)}")
        except pd.errors.EmptyDataError:
            self.bank_status_label.config(text="Excel file is empty or corrupted", fg=self.app.error_color)
            logging.error("Excel file is empty or corrupted")
            messagebox.showerror("Error", "Excel file is empty or corrupted")
        except Exception as e:
            self.bank_status_label.config(text=f"Error: {str(e)}", fg=self.app.error_color)
            logging.error(f"Unexpected error: {str(e)}")
            messagebox.showerror("Error", f"An unexpected error occurred: {str(e)}")

    def generate_word_letter(self, case, output_path):
        # 1. Look in the folder chosen by the user
        template_path = os.path.join(self.bank_template_dir, "bank.docx")

            # Add these debug lines
        print(f"DEBUG: bank_template_dir = {self.bank_template_dir}")
        print(f"DEBUG: template_path = {template_path}")
        print(f"DEBUG: os.path.exists(template_path) = {os.path.exists(template_path)}")
        print(f"DEBUG: sys.frozen = {getattr(sys, 'frozen', False)}")

        # 2. If it is not there and we are running from a PyInstaller exe,
        #    fall back to the copy that may have been embedded in the bundle.
        if not os.path.exists(template_path) and getattr(sys, "frozen", False):
            template_path = os.path.join(
                sys._MEIPASS,               # noqa:  (set by PyInstaller)
                "templates", "banks", "bank.docx"
            )
            print(f"DEBUG: Fallback template_path = {template_path}")
            print(f"DEBUG: Fallback exists = {os.path.exists(template_path)}")

        # 3. Final check – abort with a clear message if the file is still missing
        if not os.path.exists(template_path):
            err_msg = (f"Template file 'bank.docx' not found.\n\n"
                    f"Searched in:\n • {self.bank_template_dir}\n"
                    "Make sure that directory contains banks\\bank.docx.")
            self.bank_status_label.config(text=err_msg, fg=self.app.error_color)
            logging.error(f"Template file not found: {template_path}")
            messagebox.showerror("Template Missing", err_msg)
            raise FileNotFoundError(f"Template file not found: {template_path}")
        try:
            doc = Document(template_path)
        except Exception as e:
            self.bank_status_label.config(text=f"Failed to load template: {str(e)}", fg=self.app.error_color)
            logging.error(f"Failed to load template: {str(e)}")
            messagebox.showerror("Error", f"Failed to load template: {str(e)}")
            raise
        self.app.fetch_officer_details()
        replacements = {
            '{{Officer_Name}}': self.app.officer.get('OfficerName', 'Unknown Officer'),
            '{{Officer_Designation}}': self.app.officer.get('Designation', 'Unknown Designation'),
            '{{Officer_Phone}}': self.app.officer.get('Phone', 'N/A'),
            '{{Officer_Email}}': self.app.officer.get('Email', 'N/A'),
            '{{Letter_Date}}': case.get('RequestDate', datetime.now().strftime("%d-%m-%Y")),
            '{{Nodal_Officer}}': case.get('RecipientName', 'Nodal Officer'),
            '{{Bank}}': case.get('Bank', 'Unknown Bank'),
            '{{Crime_No_with_Section}}': case.get('CrimeNumber', 'N/A'),
            '{{NCRP_ID}}': case.get('NCRP_ID', 'N/A'),
            '{{Total_Amount}}': case.get('Total_Amount', 'N/A'),
            '{{Date_From}}': case.get('Date_From', 'N/A'),
            '{{Date_To}}': case.get('Date_To', 'N/A'),
        }
        logging.debug(f"Replacements for {case['Bank']}: {replacements}")
        table_inserted = False
        for paragraph in doc.paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)
            if '{{Accounts}}' in full_text and case.get('Accounts'):
                logging.debug(f"Found {{Accounts}} in paragraph")
                for run in paragraph.runs:
                    run.text = ''
                try:
                    account_table = doc.add_table(rows=len(case['Accounts']) + 1, cols=2)
                    account_table.style = 'Table Grid'
                    headers = ['Account Number', 'IFSC Code']
                    for idx, header in enumerate(headers):
                        cell = account_table.cell(0, idx)
                        cell.text = header
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                    for i, account in enumerate(case['Accounts'], start=1):
                        account_table.cell(i, 0).text = self.clean_account_number(account.get('account_no', 'N/A'))
                        account_table.cell(i, 1).text = account.get('ifsc_code', 'N/A')
                    paragraph._p.addnext(account_table._tbl)
                    table_inserted = True
                    logging.debug(f"Inserted table for {{Accounts}} with {len(case['Accounts'])} rows")
                except Exception as e:
                    logging.error(f"Failed to insert table: {str(e)}")
                    self.bank_status_label.config(text=f"Failed to insert table: {str(e)}", fg=self.app.error_color)
                    raise
            else:
                replace_placeholder_in_paragraph(paragraph, replacements)
                logging.debug(f"Post-replacement paragraph text: {paragraph.text}")
        if not table_inserted:
            logging.debug(f"No table inserted (Accounts: {len(case.get('Accounts', []))})")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, replacements)
                        logging.debug(f"Post-replacement table cell paragraph text: {paragraph.text}")
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            logging.debug(f"Saved bank letter: {output_path}")
        except Exception as e:
            self.bank_status_label.config(text=f"Failed to save letter: {str(e)}", fg=self.app.error_color)
            logging.error(f"Failed to save letter: {str(e)}")
            messagebox.showerror("Error", f"Failed to save letter to '{output_path}': {str(e)}")
            raise

    def view_letters_bank(self):
        folder_path = os.path.join(Path.home(), 'Documents', 'GeneratedLetters', 'bank')
        try:
            if os.path.exists(folder_path):
                os.startfile(folder_path)
                logging.debug(f"Opened bank letters folder: {folder_path}")
            else:
                messagebox.showwarning("Warning", "No letters generated or folder not found")
                logging.warning(f"Bank letters folder not found: {folder_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open folder: {str(e)}")
            logging.error(f"Failed to open bank letters folder: {str(e)}")